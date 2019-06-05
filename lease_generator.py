###########################################
# Dashboard for generating leases
###########################################
# Version 1.1
# Author: James Stevick
# Date: 5/30/2019
# Python 3.0
###########################################

#######################################   
# IMPORTS
####################################### 
import docx
import os
import datetime as dt
from os.path import expanduser
from tkinter import *
from tkinter import messagebox
from tkinter import filedialog

#######################################   
# DASHBOARD
####################################### 
class Dashboard_GUI(Tk):
    
    def __init__(self):

        # Setup
        Tk.__init__(self)
        self.iconbitmap('me.ico')
        self.export_path = expanduser("~") + '\\Desktop'
        self.me = expanduser("~").split("\\")[-1]
        self.template = docx.Document('example.docx')
        
        # Document Information
        self.properties = ['1636 Walnut Street','1722 Walnut Street','Carriage House','Vine Street Villas','Mango Cove','Coral Cove','Glen Cove']
        self.property_info = {
            '1636 Walnut Street':{'addr1':'1636 Walnut Street','addr2':'Berkeley, CA 94709'},
            '1722 Walnut Street':{'addr1':'1722 Walnut Street','addr2':'Berkeley CA, 94709'},
            'Carriage House':{'addr1':'1716 Rose Street','addr2':'Berkeley, CA 94703'},
            'Vine Street Villas':{'addr1':'1446 MLK Jr Way','addr2':'Berkeley, CA 94709'},
            'Mango Cove':{'addr1':'301 NE 6th Ave','addr2':'Delray Beach, FL 33483'},
            'Coral Cove':{'addr1':'88 S. Ocean Blvd','addr2':'Delray Beach, FL 33483'},
            'Glen Cove':{'addr1':'100 Ocean Ave','addr2':'Kennebunkport, ME 04046'}
        }

        # Form Block
        self.form_frame = Frame(self)
        self.form_frame.grid(row=0, column=0, sticky=N + S + W + E, padx=4, pady=4)
        self.f = FormBlock(self.form_frame, self)


#######################################   
# LEASE FORM
####################################### 
class FormBlock:

    def __init__(self, master, dashboard):
        self.dashboard = dashboard
        self.master = master
        self.default_property = 'Click to Select'
        self.form_okay = False
        self.pets_allowed = False
        self.secdep_default = '500'
        self.cleanfee_default = '125'
        self.utilmonths_default = '12'

        # DEFINE ENTRY FIELDS
        self.building_label = Label(master, text="Property:   ")
        self.building_label.grid(row=0, column=0, sticky=E)

        self.build = StringVar()
        self.build.trace("w", self.get_property)
        self.build.set(self.default_property)
        self.building_entry = OptionMenu(master, self.build, *self.dashboard.properties)
        self.building_entry.grid(row=0, column=1, padx=4, pady=4)

        self.unit_label = Label(master, text="Unit:   ")
        self.unit_label.grid(row=1, column=0, sticky=E)

        self.unit_entry = Entry(master, width=25)
        self.unit_entry.grid(row=1, column=1, padx=4, pady=4)

        self.name_label = Label(master, text="Tenant Name:   ")
        self.name_label.grid(row=3, column=0, sticky=E)

        self.name_entry = Entry(master, width=25)
        self.name_entry.grid(row=3, column=1, padx=4, pady=4)

        self.second_name_label = Label(master, text="Second Tenant:   ")
        self.second_name_label.grid(row=4, column=0, sticky=E)

        self.second_name_entry = Entry(master, width=25)
        self.second_name_entry.grid(row=4, column=1, padx=4, pady=4)

        self.email_label = Label(master, text="E-mail:   ")
        self.email_label.grid(row=5, column=0, sticky=E)

        self.email_entry = Entry(master, width=25)
        self.email_entry.grid(row=5, column=1, padx=4, pady=4)

        self.phone_label = Label(master, text="Phone #:   ")
        self.phone_label.grid(row=6, column=0, sticky=E)

        self.phone_entry = Entry(master, width=25)
        self.phone_entry.grid(row=6, column=1, padx=4, pady=4)

        self.datein_label = Label(master, text="Date In:   ")
        self.datein_label.grid(row=7, column=0, sticky=E)

        self.datein_entry = Entry(master, width=25)
        self.datein_entry.grid(row=7, column=1, padx=4, pady=4)

        self.dateout_label = Label(master, text="Date Out:   ")
        self.dateout_label.grid(row=8, column=0, sticky=E)

        self.dateout_entry = Entry(master, width=25)
        self.dateout_entry.grid(row=8, column=1, padx=4, pady=4)

        self.rent_label = Label(master, text="Rent:   ")
        self.rent_label.grid(row=9, column=0, sticky=E)

        self.rent_entry = Entry(master, width=25)
        self.rent_entry.grid(row=9, column=1, padx=4, pady=4)

        self.secdep_label = Label(master, text="Security Deposit:   ")
        self.secdep_label.grid(row=10, column=0, sticky=E)

        self.secdep_entry = Entry(master, width=25)
        self.secdep_entry.grid(row=10, column=1, padx=4, pady=4)
        self.secdep_entry.insert(0,self.secdep_default)

        self.cleanfee_label = Label(master, text="Cleaning Fee:   ")
        self.cleanfee_label.grid(row=11, column=0, sticky=E)

        self.cleanfee_entry = Entry(master, width=25)
        self.cleanfee_entry.grid(row=11, column=1, padx=4, pady=4)     
        self.cleanfee_entry.insert(0,self.cleanfee_default)

        self.utilmonths_label = Label(master, text="Months Utils Paid:   ")
        self.utilmonths_label.grid(row=12, column=0, sticky=E)

        self.utilmonths_entry = Entry(master, width=25)
        self.utilmonths_entry.grid(row=12, column=1, padx=4, pady=4)   
        self.utilmonths_entry.insert(0,self.utilmonths_default)

        self.pets_var = BooleanVar()
        self.check_button_pets = Checkbutton(master, text="Pets Allowed", variable=self.pets_var, command=self.add_pet_fields)
        self.check_button_pets.grid(row=13, column=0, padx=4, pady=4)

        # Enter Button
        self.enter_button = Button(master, text="Generate Lease", width=5, height = 2, font = 1, bg="green", command=self.generate_lease)
        self.enter_button.grid(row=18, column=0, columnspan = 2, sticky=N+E+S+W, padx=10, pady=10)       


    # VALIDATES AND CREATES DATA PACKET
    def form_check(self):

        self.form_okay = False
        error_list = []

        # Property
        self.property = self.get_property()
        if not self.property in self.dashboard.properties:
            messagebox.showerror("Error", "Property Error: Must be in list")
            return

        # Unit #
        self.unit = self.get_unit()
        if self.unit == '':
            messagebox.showerror("Error", "Unit # Not Entered")
            return
        # elif self.unit.isdigit()==False:
        #     messagebox.showerror("Error", "Unit Error: Must be a number")
        #     return  

        # Name of Tenant
        self.name = self.get_name()
        self.name = self.name.title()
        if self.name == '':
            messagebox.showerror("Error", "Name Not Entered")
            return
        elif self.name.replace(' ','').isalpha() == False:
            messagebox.showerror("Error", "Name Error: Incorrect Format")
            return

        # Name of Second Tenant
        self.second_name = self.get_second_name()
        self.second_name = self.second_name.title()
        if self.second_name.replace(' ','').isalpha() == False:
            if self.second_name == '':
                for i in range(len(self.name)):
                    self.second_name += '_'
            else:
                messagebox.showerror("Error", "Second Name Error: Incorrect Format")
                return

        # Email
        self.email = self.get_email()
        if self.email == '':
            messagebox.showerror("Error", "Email Not Entered")
            return
        elif "@" not in self.email:
            messagebox.showerror("Error", "Email Error: Incorrect Format")
            return

        # Phone Number
        self.phone = self.get_phone()
        # self.phone = self.phone.replace('(','').replace(')','').replace('-','')
        # if len(self.phone) != 10:
        #     messagebox.showerror("Error", "Phone Number Error: Must be 10 digits")
        #     return
        # elif self.phone.isdigit()==False:
        #     messagebox.showerror("Error", "Phone Number Error: Incorrect Format")
        #     return
        # self.phone = '(' + self.phone[:3]+') '+self.phone[3:6]+' - '+self.phone[7:] 

        # Date In
        self.datein = self.get_datein()
        if self.datein == '':
            messagebox.showerror("Error", "Date In Not Entered")
            return

        # Date Out
        self.dateout = self.get_dateout()
        if self.dateout == '':
            messagebox.showerror("Error", "Date Out Not Entered")
            return

        # Rent
        self.rent = self.get_rent()
        if self.rent == '':
            messagebox.showerror("Error", "Rent Not Entered")
            return
        elif self.rent.isdigit() == False:
            messagebox.showerror("Error", "Rent Error: Incorrect Format")
            return

        # Security Deposit
        self.secdep = self.get_secdep()
        if self.secdep == '':
            messagebox.showerror("Error", "Security Deposit Not Entered")
            return
        elif self.secdep.isdigit() == False:
            messagebox.showerror("Error", "Security Deposit Error: Incorrect Format")
            return

        # Cleaning Fee
        self.cleanfee = self.get_cleanfee()
        if self.cleanfee == '':
            messagebox.showerror("Error", "Cleaning Fee Not Entered")
            return
        elif self.cleanfee.isdigit() == False:
            messagebox.showerror("Error", "Cleaning Fee Error: Incorrect Format")
            return

        # Months of Paid Utilities
        self.utilmonths = self.get_utilmonths()
        if self.utilmonths == '':
            messagebox.showerror("Error", "Months of Paid Utilities Not Entered")
            return
        elif self.utilmonths.isdigit() == False:
            messagebox.showerror("Error", "Months of Paid Utilities Error: Incorrect Format")
            return

        self.today = dt.datetime.today().strftime("%m/%d/%Y")

        if self.pets_allowed:

            self.dogspecies = self.get_dogspecies()
            if len(self.dogspecies) < 20:
                for i in range(20-len(self.dogspecies)):
                    self.dogspecies += '_'
            self.dogbreed = self.get_dogbreed()
            if len(self.dogbreed) < 20:
                for i in range(20-len(self.dogbreed)):
                    self.dogbreed += '_'
            self.dogweight = self.get_dogweight()
            if len(self.dogweight) < 20:
                for i in range(20-len(self.dogweight)):
                    self.dogweight += '_'
            self.dogname = self.get_dogname()
            if len(self.dogname) < 20:
                for i in range(20-len(self.dogname)):
                    self.dogname += '_'

        self.form_okay = True

    # GET FUNCTIONS FOR FORM
    def get_property(self, *args):
        self.property = self.build.get()
        return self.property

    def get_unit(self):
        self.unit = self.unit_entry.get()
        return self.unit

    def get_name(self):
        self.name = self.name_entry.get()
        return self.name

    def get_second_name(self):
        self.second_name = self.second_name_entry.get()
        return self.second_name

    def get_email(self):
        self.email = self.email_entry.get()
        return self.email

    def get_phone(self):
        self.phone = self.phone_entry.get()
        return self.phone

    def get_datein(self):
        self.datein = self.datein_entry.get()
        return self.datein

    def get_dateout(self):
        self.dateout = self.dateout_entry.get()
        return self.dateout

    def get_rent(self):
        self.rent = self.rent_entry.get()
        return self.rent

    def get_secdep(self):
        self.secdep = self.secdep_entry.get()
        return self.secdep

    def get_cleanfee(self):
        self.cleanfee = self.cleanfee_entry.get()
        return self.cleanfee

    def get_utilmonths(self):
        self.utilmonths = self.utilmonths_entry.get()
        return self.utilmonths

    def add_pet_fields(self):

        self.pets_allowed = self.pets_var.get()

        if self.pets_allowed:
            self.dogspecies_label = Label(self.master, text="Dog Species:   ")
            self.dogspecies_label.grid(row=14, column=0, sticky=E)

            self.dogspecies_entry = Entry(self.master, width=25)
            self.dogspecies_entry.grid(row=14, column=1, padx=4, pady=4)   

            self.dogbreed_label = Label(self.master, text="Dog Breed:   ")
            self.dogbreed_label.grid(row=15, column=0, sticky=E)

            self.dogbreed_entry = Entry(self.master, width=25)
            self.dogbreed_entry.grid(row=15, column=1, padx=4, pady=4)   

            self.dogweight_label = Label(self.master, text="Dog Weight:   ")
            self.dogweight_label.grid(row=16, column=0, sticky=E)

            self.dogweight_entry = Entry(self.master, width=25)
            self.dogweight_entry.grid(row=16, column=1, padx=4, pady=4)   

            self.dogname_label = Label(self.master, text="Dog Name:   ")
            self.dogname_label.grid(row=17, column=0, sticky=E)

            self.dogname_entry = Entry(self.master, width=25)
            self.dogname_entry.grid(row=17, column=1, padx=4, pady=4)   

        else:
            self.dogspecies_label.grid_forget()
            self.dogspecies_entry.grid_forget() 
            self.dogbreed_label.grid_forget()
            self.dogbreed_entry.grid_forget() 
            self.dogweight_label.grid_forget()
            self.dogweight_entry.grid_forget() 
            self.dogname_label.grid_forget()
            self.dogname_entry.grid_forget()
            self.dogname_entry.grid_forget()

    def get_dogspecies(self):
        self.dogspecies = self.dogspecies_entry.get()
        return self.dogspecies

    def get_dogbreed(self):
        self.dogbreed = self.dogbreed_entry.get()
        return self.dogbreed

    def get_dogweight(self):
        self.dogweight = self.dogweight_entry.get()
        return self.dogweight

    def get_dogname(self):
        self.dogname = self.dogname_entry.get()
        return self.dogname


    # LEASE GENERATOR FUNCTION
    def generate_lease(self):

        self.form_check()
        if self.form_okay == False:
            return

        info_dict = {
            'TODAY_DATE':self.today,
            'TENANT_NAME':self.name,
            'OPT_TENANT_2':self.second_name,
            'TENANT_CELL':self.phone,
            'TENANT_EMAIL':self.email,
            'START_DATE':self.datein,
            'END_DATE':self.dateout,
            'SEC_DEP':self.secdep,
            'CLEANING_FEE':self.cleanfee,
            'UTIL_MONTHS':self.utilmonths,
            'MONTHLY_RENT':self.rent,
            'APT':self.unit,
            'PROPERTY_NAME':self.property,
            'ADDR_LINE_1':self.dashboard.property_info[self.property]['addr1'],
            'ADDR_LINE_2':self.dashboard.property_info[self.property]['addr2']
        }
        
        if self.pets_allowed:
            info_dict['DOG_SPECIES'] = self.dogspecies
            info_dict['DOG_BREED'] = self.dogbreed
            info_dict['DOG_WEIGHT'] = self.dogweight
            info_dict['DOG_NAME'] = self.dogname

        self.lease = self.dashboard.template

        # Handle pet clauses
        if self.pets_allowed:
            for para in list(range(39,41))[::-1]:
                p = self.lease.paragraphs[para]._element
                p.getparent().remove(p)
                p._p = p._element = None
        else:
            for para in list(range(65,71))[::-1]:
                p = self.lease.paragraphs[para]._element
                p.getparent().remove(p)
                p._p = p._element = None

        # Edit lease main text
        for idx1,para in enumerate(self.lease.paragraphs):
            for idx2,run in enumerate(para.runs):
                for idx3,char in enumerate(run.text):
                    if char == '[':
                        idx4 = idx3+1
                        field = ''
                        text_end = False
                        while not text_end:
                            if run.text[idx4] == ']':
                                text_end = True
                            else:
                                field += run.text[idx4]
                                idx4 += 1
                        style_tmp = run.style
                        run.text = run.text.replace('['+field+']',info_dict[field])
                        run.style = style_tmp

        # Add lease header
        run = self.lease.sections[0].header.paragraphs[0]    
        fields = []
        for idx5,char in enumerate(run.text):
            if char == '[':
                idx6 = idx5+1
                field = ''
                text_end = False
                while not text_end:
                    if run.text[idx6] == ']':
                        text_end = True
                    else:
                        field += run.text[idx6]
                        idx6 += 1
                fields.append(field)
        for field in fields:
            style_tmp = run.style
            run.text = run.text.replace('['+field+']',info_dict[field])
            run.style = style_tmp

        date = self.today.split('/')
        self.export_name = date[2]+'-'+date[0]+'-'+date[1]+'-'+self.property.replace(' ','-')+'-'+self.unit+'-'+self.name.split(' ')[-1]+'-For-Signing.docx'
        self.lease.save(self.export_name)

        self.build.set(self.default_property)
        self.unit_entry.delete(0, END)
        self.name_entry.delete(0, END)
        self.second_name_entry.delete(0, END)
        self.email_entry.delete(0, END)
        self.phone_entry.delete(0, END)
        self.datein_entry.delete(0, END)
        self.dateout_entry.delete(0, END)
        self.rent_entry.delete(0, END)
        self.secdep_entry.delete(0, END)
        self.cleanfee_entry.delete(0, END)
        self.utilmonths_entry.delete(0, END)

        if self.pets_allowed:
            self.dogspecies_entry.delete(0, END)
            self.dogbreed_entry.delete(0, END)
            self.dogweight_entry.delete(0, END)
            self.dogname_entry.delete(0, END)

        self.secdep_entry.insert(0, self.secdep_default)
        self.cleanfee_entry.insert(0, self.cleanfee_default)
        self.utilmonths_entry.insert(0, self.utilmonths_default)


#######################################   
# GUI MAIN
####################################### 
if __name__ == "__main__":
    window = Dashboard_GUI()
    window.wm_title("Lease Generator Beta")
    window.resizable(1, 1)
    window.mainloop()