import docx
doc = docx.Document('example.docx')

properties = []
properties.append({'name':'1636 Walnut Street','addr1':'1636 Walnut Street','addr2':'Berkeley, CA 94709'})
properties.append({'name':'1722 Walnut Street','addr1':'1722 Walnut Street','addr2':'Berkeley CA, 94709'})
properties.append({'name':'Carriage House','addr1':'1716 Rose Street','addr2':'Berkeley, CA 94703'})
properties.append({'name':'Vine Street Villas','addr1':'1446 MLK Jr Way','addr2':'Berkeley, CA 94709'})
properties.append({'name':'Mango Cove','addr1':'301 NE 6th Ave','addr2':'Delray Beach, FL 33483'})
properties.append({'name':'Coral Cove','addr1':'88 S. Ocean Blvd','addr2':'Delray Beach, FL 33483'})
properties.append({'name':'Glen Cove','addr1':'100 Ocean Ave','addr2':'Kennebunkport, ME 04046'})
PROP = 0

# PROPERTY_NAME = properties[PROP]['name']
# ADDR_LINE_1 = properties[PROP]['addr1']
# ADDR_LINE_2 = properties[PROP]['addr2']

# style_name = doc.paragraphs[0].runs[0].style
# style_addr1 = doc.paragraphs[1].runs[0].style
# style_addr2 = doc.paragraphs[2].runs[0].style

# doc.paragraphs[0].runs[0].text = properties[0]['name']
# doc.paragraphs[1].runs[0].text = properties[0]['addr1']
# doc.paragraphs[2].runs[0].text = properties[0]['addr2']

# doc.paragraphs[0].runs[0].style = style_name
# doc.paragraphs[1].runs[0].style = style_addr1
# doc.paragraphs[2].runs[0].style = style_addr2

info_dict = {
    'TODAY_DATE':'05/29/2019',
    'TENANT_NAME':'Jake Johnson',
    'OPT_TENANT_2':'Betsy Meloncamp',
    'TENANT_CELL':'510-501-7466',
    'TENANT_EMAIL':'jelloshot@hotmail.com',
    'START_DATE':'09/29/2019',
    'END_DATE':'12/29/2019',
    'SEC_DEP':'500',
    'CLEANING_FEE':'200',
    'UTIL_MONTHS':'8',
    'MONTHLY_RENT':'2900',
    'APT':'5',
    'PROPERTY_NAME':properties[PROP]['name'],
    'ADDR_LINE_1':properties[PROP]['addr1'],
    'ADDR_LINE_2':properties[PROP]['addr2'],
    'DOG_SPECIES':'duck',
    'DOG_BREED':'sauce',
    'DOG_WEIGHT':'150',
    'DOG_NAME':'Larry'
}


for idx1,para in enumerate(doc.paragraphs):
    for idx2,run in enumerate(para.runs):
        for idx3,char in enumerate(run.text):
            if char == '[':
                idx4 = idx3+1
                field = ''
                text_end = False
                while not text_end:
                    if run.text[idx4] == ']':
                        text_end = True
                    else:
                        field += run.text[idx4]
                        idx4 += 1
                print(field)
                style_tmp = run.style
                run.text = run.text.replace('['+field+']',info_dict[field])
                run.style = style_tmp
                
run = doc.sections[0].header.paragraphs[0]    
fields = []
for idx5,char in enumerate(run.text):
    if char == '[':
        idx6 = idx5+1
        field = ''
        text_end = False
        while not text_end:
            if run.text[idx6] == ']':
                text_end = True
            else:
                field += run.text[idx6]
                idx6 += 1
        print(field)
        fields.append(field)
for field in fields:
    style_tmp = run.style
    run.text = run.text.replace('['+field+']',info_dict[field])
    run.style = style_tmp


for para in list(range(65,71))[::-1]:
    p = doc.paragraphs[para]._element
    p.getparent().remove(p)
    p._p = p._element = None


doc.save('updated.docx')